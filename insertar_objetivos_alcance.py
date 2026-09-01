from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from copy import deepcopy

src = r"C:\Users\Ulises\Downloads\Seminario Integracion - Belazquez y Serrano.docx"
out = r"C:\Users\Ulises\OneDrive\Desktop\Seminario Integracion - Sistema Villafañe\Seminario Integracion - Villafañe Wifi - actualizado.docx"

doc = Document(src)

def insert_paragraph_after(paragraph, text="", style=None, bold_prefix=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if bold_prefix and text.startswith(bold_prefix):
        r1 = new_para.add_run(bold_prefix)
        r1.bold = True
        new_para.add_run(text[len(bold_prefix):])
    else:
        new_para.add_run(text)
    return new_para

# Insert immediately before the existing “Ciclo de vida” heading.
anchor = next(p for p in doc.paragraphs if p.text.strip() == "Ciclo de vida")

items = [
    ("Objetivos", "Heading 1"),
    ("Objetivo general", "Heading 1"),
    ("Desarrollar un sistema de información integral para Villafañe Wifi que centralice la gestión de clientes, servicios, cuotas y pagos, y que facilite la atención inicial mediante un bot de WhatsApp, con el fin de reducir la carga operativa, mejorar la trazabilidad de la información y brindar respuestas más ágiles a los clientes.", None),
    ("Objetivos específicos", "Heading 1"),
    ("• Centralizar en un panel web la información de clientes, sus servicios o conexiones, planes contratados, cuotas y estado de cuenta.", None),
    ("• Registrar, consultar y actualizar los datos necesarios para la administración de clientes y servicios, contemplando que un cliente pueda tener más de una conexión.", None),
    ("• Implementar un canal de atención inicial por WhatsApp que permita identificar la intención del cliente, responder consultas frecuentes y registrar reclamos.", None),
    ("• Recibir comprobantes de pago, extraer sus datos mediante OCR y facilitar la validación de comprobantes, evitando duplicaciones.", None),
    ("• Registrar los pagos aprobados y asociarlos con la cuota correspondiente para mantener actualizada la cuenta corriente.", None),
    ("• Organizar los tickets de soporte por orden de llegada y asignarlos al primer empleado disponible para su atención.", None),
    ("• Proporcionar información consistente entre el bot de WhatsApp y el panel web mediante un backend y una base de datos centralizados.", None),
    ("Alcance y exclusiones", "Heading 1"),
    ("Alcance incluido", "Heading 1"),
    ("El proyecto contempla el análisis, diseño y desarrollo incremental de un sistema compuesto por un panel web interno y un bot de WhatsApp integrados mediante una API central. El alcance funcional inicial incluye la gestión de clientes y servicios, la administración de planes y cuotas, la consulta de cuenta corriente, la recepción y validación de comprobantes, el registro de pagos, la creación y asignación de tickets de soporte y la atención inicial de conversaciones por WhatsApp. También incluye la autenticación y diferenciación de usuarios internos mediante un esquema de generalización/especialización entre Usuario, Empleado y Administrador, además de los reportes y consultas necesarios para el seguimiento operativo.", None),
    ("Exclusiones y trabajo futuro", "Heading 1"),
    ("Queda fuera del alcance principal la integración directa con equipos o servicios de red MikroTik para realizar cortes, activaciones o cambios automáticos del servicio. Durante esta etapa, el sistema registrará el estado de la cuenta y podrá generar la información necesaria para que el personal tome la decisión correspondiente, pero cualquier acción sobre la infraestructura de red se efectuará por fuera del sistema. La integración con MikroTik podrá evaluarse como una ampliación futura, una vez validado el funcionamiento del núcleo administrativo y de atención.", None),
    ("Asimismo, la primera versión no contempla una aplicación móvil nativa para clientes, un sistema de facturación fiscal, pagos parciales, liquidación automática de medios de pago ni el reemplazo inmediato de la librería no oficial de WhatsApp por la API oficial de Meta. Estos aspectos podrán incorporarse en iteraciones posteriores según las prioridades de la empresa, la disponibilidad de servicios externos y los resultados de las pruebas.", None),
]

# Insert in source order; each paragraph is inserted immediately before the
# anchor, so the anchor is moved forward after every insertion.
for text, style in items:
    p = OxmlElement("w:p")
    anchor._p.addprevious(p)
    para = Paragraph(p, anchor._parent)
    if style:
        para.style = style
    para.add_run(text)

doc.save(out)
print(out)
