from docx import Document
from pathlib import Path
import sys

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

path = Path(r"C:\Users\Ulises\Downloads\Seminario Integracion - Belazquez y Serrano(2).docx")
doc = Document(path)

items = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        items.append((f"P{i}", text))
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        text = " | ".join(cell.text.strip().replace("\n", " / ") for cell in row.cells)
        if text.strip(" |"):
            items.append((f"T{ti}R{ri}", text))

needles = ("RF-13", "RF 13", "transfer", "escalad", "RF-12", "RF 12")
for location, text in items:
    if any(n.lower() in text.lower() for n in needles):
        print(f"{location}: {text}")

print(f"PARAGRAPHS={len(doc.paragraphs)} TABLES={len(doc.tables)}")

print("\nREQUIREMENTS")
for location, text in items:
    if text.lstrip().startswith("RF-"):
        print(f"{location}: {text}")

print("\nBUSINESS CONTEXT")
business_needles = (
    "vencimiento", "comprobante", "pago", "ticket", "reclamo", "plan",
    "cliente", "whatsapp", "cuenta corriente", "usuario", "empleado",
    "administrador", "ocr", "dni", "mac", "ip", "servicio",
)
for location, text in items:
    lower = text.lower()
    if any(n in lower for n in business_needles) and not text.lstrip().startswith("RF-"):
        print(f"{location}: {text}")
