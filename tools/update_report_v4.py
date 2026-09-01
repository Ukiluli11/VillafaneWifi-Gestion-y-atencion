from copy import deepcopy
from pathlib import Path
import re

from docx import Document

SOURCE = Path(r"C:\Users\Ulises\Downloads\Seminario Integracion - Belazquez y Serrano(2).docx")
OUTPUT = Path(r"C:\Users\Ulises\OneDrive\Desktop\Seminario Integracion - Sistema Villafañe\outputs\revision_modelo_v4\Seminario Integracion - Belazquez y Serrano - V4.docx")


def replace_paragraph_text(paragraph, old, new):
    if old not in paragraph.text:
        return False
    full = paragraph.text.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = full
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(full)
    return True


def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    for extra in cell.paragraphs[1:]:
        for run in extra.runs:
            run.text = ""


doc = Document(SOURCE)

# Ajustes de redacción general para reflejar persistencia de mensajes,
# pagos que cancelan varias cuotas y atención por administradores.
paragraph_replacements = {
    "identificar la intención del cliente, responder consultas frecuentes y registrar reclamos.":
        "identificar la intención del cliente, responder consultas frecuentes, registrar reclamos y conservar el historial de mensajes de cada conversación.",
    "Registrar los pagos aprobados y asociarlos con la cuota correspondiente para mantener actualizada la cuenta corriente.":
        "Registrar los pagos aprobados y asociarlos con una o varias cuotas o servicios correspondientes para mantener actualizada la cuenta corriente.",
    "la creación y asignación de tickets de soporte y la atención inicial de conversaciones por WhatsApp.":
        "la creación y asignación de tickets de soporte, la atención inicial de conversaciones por WhatsApp y la conservación de su historial de mensajes.",
    "el sistema podrá identificar la intención del cliente, por ejemplo, si desea realizar un pago, reportar un problema con el servicio, o consultar el estado de su cuenta, y guiarlo en consecuencia.":
        "el sistema podrá identificar la intención del cliente, por ejemplo, si desea realizar un pago, reportar un problema con el servicio, o consultar el estado de su cuenta, y guiarlo en consecuencia. Cada mensaje enviado o recibido se conservará como parte del historial de la conversación, de modo que el bot y los usuarios internos autorizados puedan consultar el contexto durante la atención.",
}
for paragraph in doc.paragraphs:
    for old, new in paragraph_replacements.items():
        replace_paragraph_text(paragraph, old, new)

# Renumerar requerimientos funcionales posteriores al RF-12.
for table in doc.tables:
    for row in table.rows:
        code = row.cells[0].text.strip() if row.cells else ""
        match = re.fullmatch(r"RF-(\d{2})", code)
        if match and int(match.group(1)) >= 13:
            set_cell_text(row.cells[0], f"RF-{int(match.group(1)) + 1:02d}")

# Incorporar el nuevo RF-13 al final del módulo de WhatsApp.
bot_table = next(
    table for table in doc.tables
    if any(row.cells and row.cells[0].text.strip() == "RF-12" for row in table.rows)
)
new_tr = deepcopy(bot_table.rows[-1]._tr)
bot_table._tbl.append(new_tr)
new_row = bot_table.rows[-1]
new_values = [
    "RF-13",
    "Registro del historial de mensajes",
    "El sistema debe almacenar cada mensaje enviado o recibido en una conversación de WhatsApp, registrando fecha y hora, contenido o archivo adjunto, tipo de mensaje y emisor (cliente, bot o usuario interno), de modo que el historial pueda consultarse durante el escalado y desde el panel web.",
    "Media",
]
for cell, value in zip(new_row.cells, new_values):
    set_cell_text(cell, value)

# Ajustes puntuales de requerimientos afectados por las decisiones del modelo.
requirement_updates = {
    "RF-12": (
        "Escalado de conversación a un usuario interno",
        "El sistema debe permitir que un empleado o administrador autorizado tome el control de una conversación en curso cuando el bot no pueda resolver la consulta del cliente.",
    ),
    "RF-18": (
        "Confirmación o rechazo de comprobante",
        "El sistema debe permitir a un empleado o administrador autorizado confirmar o rechazar un comprobante desde la bandeja de conciliación.",
    ),
    "RF-19": (
        "Actualización automática de cuenta corriente",
        "Al confirmar un comprobante, el sistema debe registrar el pago, asociarlo con una o varias cuotas o servicios y actualizar automáticamente la cuenta corriente y el historial de pagos del cliente.",
    ),
    "RF-22": (
        "Asignación de ticket a un responsable",
        "El sistema debe permitir asignar un ticket a un empleado o administrador autorizado, respetando el orden de llegada y la disponibilidad del personal.",
    ),
    "RF-23": (
        "Cambio de estado de ticket",
        "El sistema debe permitir que el usuario interno responsable cambie el estado de un ticket (pendiente, en proceso, resuelto).",
    ),
    "RF-24": (
        "Comentarios internos en ticket",
        "El sistema debe permitir que empleados y administradores autorizados agreguen notas o comentarios internos a un ticket.",
    ),
    "RF-30": (
        "Niveles de acceso por tipo de usuario y área",
        "El sistema debe permitir asignar niveles de acceso según el tipo de usuario (empleado o administrador) y, para los empleados, según su área de trabajo (administración, soporte técnico o atención al cliente).",
    ),
}
for table in doc.tables:
    for row in table.rows:
        if not row.cells:
            continue
        code = row.cells[0].text.strip()
        if code in requirement_updates:
            name, description = requirement_updates[code]
            set_cell_text(row.cells[1], name)
            set_cell_text(row.cells[2], description)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
