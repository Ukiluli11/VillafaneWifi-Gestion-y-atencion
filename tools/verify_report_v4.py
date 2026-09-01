from pathlib import Path
from zipfile import ZipFile
import re
import sys

from docx import Document

sys.stdout.reconfigure(encoding="utf-8", errors="replace")
path = Path(r"C:\Users\Ulises\OneDrive\Desktop\Seminario Integracion - Sistema Villafañe\outputs\revision_modelo_v4\Seminario Integracion - Belazquez y Serrano - V4.docx")

with ZipFile(path) as zf:
    bad = zf.testzip()
    print(f"ZIP_OK={bad is None}")

doc = Document(path)
requirements = []
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        code = row.cells[0].text.strip() if row.cells else ""
        if re.fullmatch(r"RF-\d{2}", code):
            requirements.append((code, ti, ri, row.cells[1].text.strip(), row.cells[2].text.strip()))

codes = [code for code, *_ in requirements]
expected = [f"RF-{i:02d}" for i in range(1, 31)]
print(f"RF_COUNT={len(codes)} CONTINUOUS={codes == expected} UNIQUE={len(set(codes)) == len(codes)}")
for item in requirements:
    if item[0] in {"RF-12", "RF-13", "RF-14", "RF-18", "RF-19", "RF-22", "RF-24", "RF-29", "RF-30"}:
        print(" | ".join(map(str, item)))

all_text = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
    " | ".join(cell.text for cell in row.cells)
    for table in doc.tables for row in table.rows
)
checks = {
    "message_history": "historial de mensajes" in all_text.lower(),
    "multi_quota": "una o varias cuotas" in all_text.lower(),
    "rf30": "RF-30" in all_text,
    "old_rf29_access": "RF-29 | Niveles de acceso" in all_text,
}
print("CHECKS", checks)
print(f"PARAGRAPHS={len(doc.paragraphs)} TABLES={len(doc.tables)}")
