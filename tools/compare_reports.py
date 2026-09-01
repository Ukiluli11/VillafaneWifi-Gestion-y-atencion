from __future__ import annotations

import difflib
import re
import sys
from collections import OrderedDict
from pathlib import Path

from docx import Document


def normalize(text: str) -> str:
    return " ".join(text.replace("\u00a0", " ").split())


def extract(path: Path):
    doc = Document(path)
    paragraphs = [normalize(p.text) for p in doc.paragraphs if normalize(p.text)]
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            values = [normalize(cell.text) for cell in row.cells]
            if any(values):
                rows.append(values)
        if rows:
            tables.append(rows)
    return paragraphs, tables


def requirements(tables):
    result = OrderedDict()
    for table in tables:
        for row in table:
            joined = " | ".join(row)
            match = re.search(r"\bRF[-\s]?(\d{1,2})\b", joined, flags=re.I)
            if match:
                key = f"RF-{int(match.group(1)):02d}"
                result[key] = joined
    return result


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: compare_reports.py <attached.docx> <v4.docx>")
    source, v4 = map(Path, sys.argv[1:])
    source_p, source_t = extract(source)
    v4_p, v4_t = extract(v4)
    source_rf, v4_rf = requirements(source_t), requirements(v4_t)

    print(f"SOURCE paragraphs={len(source_p)} tables={len(source_t)} rf={len(source_rf)}")
    print(f"V4 paragraphs={len(v4_p)} tables={len(v4_t)} rf={len(v4_rf)}")
    print("\n=== REQUIREMENTS ONLY IN SOURCE ===")
    for key in source_rf.keys() - v4_rf.keys(): print(key, source_rf[key])
    print("\n=== REQUIREMENTS ONLY IN V4 ===")
    for key in v4_rf.keys() - source_rf.keys(): print(key, v4_rf[key])
    print("\n=== CHANGED REQUIREMENTS ===")
    changed = 0
    for key in source_rf.keys() & v4_rf.keys():
        if source_rf[key] != v4_rf[key]:
            changed += 1
            print(f"\n{key}\nSOURCE: {source_rf[key]}\nV4:     {v4_rf[key]}")
    print(f"changed_requirements={changed}")

    print("\n=== ADDED/REMOVED PARAGRAPH BLOCKS (context-independent) ===")
    sm = difflib.SequenceMatcher(a=source_p, b=v4_p, autojunk=False)
    change_count = 0
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            continue
        change_count += 1
        print(f"\n[{tag}] source {i1 + 1}:{i2} -> V4 {j1 + 1}:{j2}")
        for item in source_p[i1:i2]: print("-", item)
        for item in v4_p[j1:j2]: print("+", item)
    print(f"paragraph_change_blocks={change_count}")

    print("\n=== TABLE SHAPE ===")
    print("SOURCE", [len(t) for t in source_t])
    print("V4", [len(t) for t in v4_t])


if __name__ == "__main__":
    main()
