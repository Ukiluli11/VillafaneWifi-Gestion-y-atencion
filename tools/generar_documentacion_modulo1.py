"""Genera los tres documentos formales de cierre del Módulo 1."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


RAIZ = Path(__file__).resolve().parents[1]
SALIDA = RAIZ / "docs" / "entregables_modulo_1"
LOGO = RAIZ / "backend" / "static" / "img" / "logo-villafane-wifi.png"

VERDE_OSCURO = "0E3A2E"
VERDE = "24936A"
VERDE_CLARO = "E7F6EF"
GRIS = "687970"
GRIS_CLARO = "F2F4F3"
ROJO = "A42E3B"
BLANCO = "FFFFFF"
NEGRO = "17251F"


def configurar_fuente(run, tamano=11, color=NEGRO, negrita=False, cursiva=False):
    """Aplica tipografía explícita para una renderización estable."""

    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(tamano)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = negrita
    run.italic = cursiva


def sombrear_celda(celda, color):
    """Asigna un color de fondo a una celda."""

    propiedades = celda._tc.get_or_add_tcPr()
    sombreado = propiedades.find(qn("w:shd"))
    if sombreado is None:
        sombreado = OxmlElement("w:shd")
        propiedades.append(sombreado)
    sombreado.set(qn("w:fill"), color)


def configurar_margenes_celda(celda, superior=90, inferior=90, inicio=120, fin=120):
    """Establece márgenes internos uniformes en unidades DXA."""

    propiedades = celda._tc.get_or_add_tcPr()
    margenes = propiedades.first_child_found_in("w:tcMar")
    if margenes is None:
        margenes = OxmlElement("w:tcMar")
        propiedades.append(margenes)
    for nombre, valor in {
        "top": superior,
        "bottom": inferior,
        "start": inicio,
        "end": fin,
    }.items():
        elemento = margenes.find(qn(f"w:{nombre}"))
        if elemento is None:
            elemento = OxmlElement(f"w:{nombre}")
            margenes.append(elemento)
        elemento.set(qn("w:w"), str(valor))
        elemento.set(qn("w:type"), "dxa")


def configurar_geometria_tabla(tabla, anchos):
    """Fija el ancho total, la grilla, la sangría y cada columna de la tabla."""

    total = sum(anchos)
    propiedades = tabla._tbl.tblPr
    ancho_tabla = propiedades.find(qn("w:tblW"))
    ancho_tabla.set(qn("w:w"), str(total))
    ancho_tabla.set(qn("w:type"), "dxa")

    sangria = propiedades.find(qn("w:tblInd"))
    if sangria is None:
        sangria = OxmlElement("w:tblInd")
        propiedades.append(sangria)
    sangria.set(qn("w:w"), "120")
    sangria.set(qn("w:type"), "dxa")

    grilla = tabla._tbl.tblGrid
    for columna in list(grilla):
        grilla.remove(columna)
    for ancho in anchos:
        columna = OxmlElement("w:gridCol")
        columna.set(qn("w:w"), str(ancho))
        grilla.append(columna)

    tabla.autofit = False
    for fila in tabla.rows:
        for indice, celda in enumerate(fila.cells):
            propiedades_celda = celda._tc.get_or_add_tcPr()
            ancho_celda = propiedades_celda.find(qn("w:tcW"))
            if ancho_celda is None:
                ancho_celda = OxmlElement("w:tcW")
                propiedades_celda.append(ancho_celda)
            ancho_celda.set(qn("w:w"), str(anchos[indice]))
            ancho_celda.set(qn("w:type"), "dxa")
            celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            configurar_margenes_celda(celda)


def repetir_encabezado(fila):
    """Marca una fila para que Word la repita al dividir la tabla."""

    propiedades = fila._tr.get_or_add_trPr()
    marcador = OxmlElement("w:tblHeader")
    marcador.set(qn("w:val"), "true")
    propiedades.append(marcador)


def agregar_numero_pagina(parrafo):
    """Inserta un campo PAGE real en el pie de página."""

    run = parrafo.add_run()
    inicio = OxmlElement("w:fldChar")
    inicio.set(qn("w:fldCharType"), "begin")
    instruccion = OxmlElement("w:instrText")
    instruccion.set(qn("xml:space"), "preserve")
    instruccion.text = " PAGE "
    separador = OxmlElement("w:fldChar")
    separador.set(qn("w:fldCharType"), "separate")
    texto = OxmlElement("w:t")
    texto.text = "1"
    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")
    for elemento in (inicio, instruccion, separador, texto, fin):
        run._r.append(elemento)
    configurar_fuente(run, tamano=9, color=GRIS)


def configurar_documento(titulo_corto, preset="compacto"):
    """Crea un documento con estilos, página, encabezado y pie coherentes."""

    documento = Document()
    seccion = documento.sections[0]
    seccion.page_width = Inches(8.5)
    seccion.page_height = Inches(11)
    seccion.top_margin = Inches(1)
    seccion.bottom_margin = Inches(1)
    seccion.left_margin = Inches(1)
    seccion.right_margin = Inches(1)
    seccion.header_distance = Inches(0.49)
    seccion.footer_distance = Inches(0.49)
    seccion.different_first_page_header_footer = True

    normal = documento.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NEGRO)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25 if preset == "compacto" else 1.10

    especificaciones = {
        "Title": (26, VERDE_OSCURO, 0, 8),
        "Subtitle": (13, GRIS, 0, 14),
        "Heading 1": (16, VERDE_OSCURO, 18, 10),
        "Heading 2": (13, VERDE, 14, 7),
        "Heading 3": (12, VERDE_OSCURO, 10, 5),
    }
    for nombre, (tamano, color, antes, despues) in especificaciones.items():
        estilo = documento.styles[nombre]
        estilo.font.name = "Aptos Display" if nombre != "Subtitle" else "Aptos"
        estilo._element.rPr.rFonts.set(qn("w:ascii"), estilo.font.name)
        estilo._element.rPr.rFonts.set(qn("w:hAnsi"), estilo.font.name)
        estilo.font.size = Pt(tamano)
        estilo.font.color.rgb = RGBColor.from_string(color)
        estilo.font.bold = nombre != "Subtitle"
        estilo.paragraph_format.space_before = Pt(antes)
        estilo.paragraph_format.space_after = Pt(despues)
        estilo.paragraph_format.keep_with_next = True

    for nombre in ("List Bullet", "List Number"):
        estilo = documento.styles[nombre]
        estilo.font.name = "Aptos"
        estilo.font.size = Pt(11)
        estilo.paragraph_format.left_indent = Inches(0.5 if preset != "compacto" else 0.375)
        estilo.paragraph_format.first_line_indent = Inches(-0.19)
        estilo.paragraph_format.space_after = Pt(4 if preset == "compacto" else 8)
        estilo.paragraph_format.line_spacing = 1.25 if preset == "compacto" else 1.167

    encabezado = seccion.header.paragraphs[0]
    encabezado.alignment = WD_ALIGN_PARAGRAPH.LEFT
    configurar_fuente(
        encabezado.add_run(f"VILLAFAÑE WIFI  |  {titulo_corto.upper()}"),
        tamano=8.5,
        color=GRIS,
        negrita=True,
    )
    encabezado.paragraph_format.space_after = Pt(4)
    borde = OxmlElement("w:pBdr")
    inferior = OxmlElement("w:bottom")
    inferior.set(qn("w:val"), "single")
    inferior.set(qn("w:sz"), "5")
    inferior.set(qn("w:color"), "BFD8CC")
    borde.append(inferior)
    encabezado._p.get_or_add_pPr().append(borde)

    pie = seccion.footer.paragraphs[0]
    pie.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    configurar_fuente(pie.add_run("Seminario de Integración  |  Página "), tamano=9, color=GRIS)
    agregar_numero_pagina(pie)
    return documento


def agregar_portada(documento, titulo, subtitulo, tipo_documento):
    """Construye una portada editorial con identidad de Villafañe Wifi."""

    parrafo_logo = documento.add_paragraph()
    parrafo_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parrafo_logo.paragraph_format.space_after = Pt(12)
    imagen = parrafo_logo.add_run().add_picture(str(LOGO), width=Inches(2.1))
    propiedades_imagen = imagen._inline.docPr
    propiedades_imagen.set("title", "Logo de Villafañe Wifi")
    propiedades_imagen.set("descr", "Identidad visual de Villafañe Wifi")

    etiqueta = documento.add_paragraph()
    etiqueta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    configurar_fuente(etiqueta.add_run(tipo_documento.upper()), tamano=10, color=VERDE, negrita=True)
    etiqueta.paragraph_format.space_after = Pt(10)

    titulo_p = documento.add_paragraph(style="Title")
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_p.add_run(titulo)

    subtitulo_p = documento.add_paragraph(style="Subtitle")
    subtitulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo_p.add_run(subtitulo)

    documento.add_paragraph()
    metadatos = [
        "Proyecto: Sistema de Gestión Integral y Atención al Cliente para Villafañe Wifi",
        "Autores: Belazquez y Serrano",
        "Carrera: Licenciatura en Sistemas de Información",
        "Materia: Seminario de Integración",
        f"Versión evaluada: Módulo 1 - {date(2026, 9, 2).strftime('%d/%m/%Y')}",
    ]
    for texto in metadatos:
        parrafo = documento.add_paragraph()
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        configurar_fuente(parrafo.add_run(texto), tamano=10, color=GRIS)
        parrafo.paragraph_format.space_after = Pt(3)
    documento.add_page_break()


def agregar_parrafo(documento, texto, negrita_inicial=None):
    """Agrega un párrafo de cuerpo con una etiqueta inicial opcional."""

    parrafo = documento.add_paragraph()
    if negrita_inicial and texto.startswith(negrita_inicial):
        configurar_fuente(parrafo.add_run(negrita_inicial), negrita=True)
        configurar_fuente(parrafo.add_run(texto[len(negrita_inicial) :]))
    else:
        configurar_fuente(parrafo.add_run(texto))
    return parrafo


def agregar_vinetas(documento, elementos):
    """Agrega una lista con viñetas reales de Word."""

    for elemento in elementos:
        parrafo = documento.add_paragraph(style="List Bullet")
        configurar_fuente(parrafo.add_run(elemento))


def agregar_tabla(documento, encabezados, filas, anchos, color_encabezado=VERDE_OSCURO):
    """Agrega una tabla tabular con geometría fija y encabezado repetible."""

    tabla = documento.add_table(rows=1, cols=len(encabezados))
    tabla.alignment = WD_TABLE_ALIGNMENT.LEFT
    tabla.style = "Table Grid"
    for indice, texto in enumerate(encabezados):
        celda = tabla.rows[0].cells[indice]
        sombrear_celda(celda, color_encabezado)
        parrafo = celda.paragraphs[0]
        configurar_fuente(parrafo.add_run(texto), tamano=9.5, color=BLANCO, negrita=True)
    repetir_encabezado(tabla.rows[0])

    for fila in filas:
        celdas = tabla.add_row().cells
        for indice, texto in enumerate(fila):
            parrafo = celdas[indice].paragraphs[0]
            configurar_fuente(parrafo.add_run(str(texto)), tamano=9.5)
    configurar_geometria_tabla(tabla, anchos)
    documento.add_paragraph().paragraph_format.space_after = Pt(2)
    return tabla


def agregar_nota(documento, etiqueta, texto, tipo="informativa"):
    """Incorpora una nota destacada de una sola columna."""

    tabla = documento.add_table(rows=1, cols=1)
    # El control de accesibilidad de Word requiere identificar la primera fila
    # incluso en estas tablas de una sola celda usadas como recuadros de aviso.
    repetir_encabezado(tabla.rows[0])
    tabla.style = "Table Grid"
    celda = tabla.cell(0, 0)
    sombrear_celda(celda, VERDE_CLARO if tipo == "informativa" else "FFF3CD")
    parrafo = celda.paragraphs[0]
    configurar_fuente(parrafo.add_run(f"{etiqueta}: "), tamano=10, color=VERDE_OSCURO, negrita=True)
    configurar_fuente(parrafo.add_run(texto), tamano=10)
    configurar_geometria_tabla(tabla, [9360])
    documento.add_paragraph().paragraph_format.space_after = Pt(2)


def guardar(documento, nombre):
    """Guarda un documento dentro del directorio final."""

    SALIDA.mkdir(parents=True, exist_ok=True)
    ruta = SALIDA / nombre
    documento.save(ruta)
    return ruta


def crear_revision_final():
    """Genera el informe de revisión y aceptación técnica del Módulo 1."""

    doc = configurar_documento("Revisión final del Módulo 1", preset="negocio")
    agregar_portada(
        doc,
        "Revisión final del Módulo 1",
        "Resultado funcional, técnico y de calidad antes del cierre académico",
        "Informe de revisión",
    )

    doc.add_heading("1. Dictamen ejecutivo", level=1)
    agregar_nota(
        doc,
        "Resultado",
        "El Módulo 1 está aprobado para demostración académica y continuidad del proyecto. "
        "No debe considerarse listo para producción hasta completar el endurecimiento de seguridad y el despliegue.",
    )
    agregar_parrafo(
        doc,
        "La revisión abarcó autenticación, permisos, clientes, planes, servicios, cuenta corriente, "
        "cuotas, pagos, base de datos, panel web, API REST, pruebas y calidad estática. El código evaluado "
        "corresponde al commit 815d13d de la rama main.",
    )

    doc.add_heading("2. Alcance revisado", level=1)
    agregar_vinetas(
        doc,
        [
            "RF-01 a RF-06: gestión comercial y cuenta corriente.",
            "RF-29 y RF-30: autenticación y permisos por subtipo y área.",
            "Persistencia PostgreSQL y migraciones de usuarios, clientes, servicios y facturación.",
            "Panel web con identidad visual de Villafañe Wifi y API REST autenticada.",
            "Pruebas automatizadas y revisión estática del código.",
        ],
    )

    doc.add_heading("3. Matriz de cumplimiento", level=1)
    filas = [
        ("RF-01", "Alta integral de cliente, teléfonos y primera conexión", "Cumple"),
        ("RF-02", "Consulta, edición y baja lógica de clientes", "Cumple"),
        ("RF-03", "Búsqueda por documento, nombre, teléfono o localidad", "Cumple"),
        ("RF-04", "Catálogo de planes y asignación a servicios", "Cumple"),
        ("RF-05", "Cuenta corriente, deuda y vencimientos", "Cumple"),
        ("RF-06", "Historial de cuotas y pagos", "Cumple"),
        ("RF-29", "Inicio y cierre de sesión", "Cumple"),
        ("RF-30", "Acceso según administrador o área del empleado", "Cumple"),
    ]
    agregar_tabla(doc, ["Req.", "Evidencia funcional", "Estado"], filas, [1050, 6660, 1650])

    doc.add_heading("4. Verificaciones ejecutadas", level=1)
    verificaciones = [
        ("Pruebas automatizadas", "50 pruebas", "Correctas"),
        ("Ruff", "Calidad estática del backend", "Sin errores"),
        ("Migraciones", "makemigrations --check --dry-run", "Sin cambios pendientes"),
        ("Django check", "Configuración funcional local", "Sin errores"),
        ("PostgreSQL", "Migraciones aplicadas", "Correcto"),
        ("Inspección visual", "Inicio de sesión y panel", "Correcto"),
    ]
    agregar_tabla(doc, ["Control", "Cobertura", "Resultado"], verificaciones, [2200, 4660, 2500])

    doc.add_heading("5. Hallazgos de la revisión", level=1)
    doc.add_heading("5.1 Fortalezas", level=2)
    agregar_vinetas(
        doc,
        [
            "Separación modular por capacidades del negocio.",
            "Lógica centralizada en servicios de aplicación orientados a objetos.",
            "Nombres propios, mensajes y documentación mayormente en español.",
            "Bajas lógicas y restricciones que preservan la historia comercial.",
            "Cuenta corriente calculada sin almacenar saldos redundantes.",
            "Misma matriz de acceso aplicada en panel y API.",
        ],
    )
    doc.add_heading("5.2 Pendientes para producción", level=2)
    pendientes = [
        ("Alta", "Configurar DEBUG=False y una DJANGO_SECRET_KEY exclusiva del entorno."),
        ("Alta", "Forzar HTTPS y habilitar cookies seguras de sesión y CSRF."),
        ("Media", "Definir HSTS cuando el dominio y certificado estén estabilizados."),
        ("Media", "Configurar copias de seguridad y restauración de PostgreSQL."),
        ("Media", "Agregar integración continua para ejecutar pruebas en cada cambio."),
        ("Baja", "Automatizar la generación mensual de cuotas mediante Celery."),
    ]
    agregar_tabla(doc, ["Prioridad", "Acción pendiente"], pendientes, [1500, 7860], color_encabezado=ROJO)

    doc.add_heading("6. Coherencia de datos y diagramas", level=1)
    agregar_parrafo(
        doc,
        "La especialización Usuario-Empleado-Administrador está implementada con clave compartida: "
        "id_usuario es simultáneamente PK y FK en los subtipos. El DER lógico debe reflejar esta decisión. "
        "Las claves principales creadas por Django se llaman id en las tablas principales; el informe académico "
        "puede usar nombres semánticos si mantiene una correspondencia explícita.",
    )
    agregar_nota(
        doc,
        "Importante",
        "La cuota no almacena estado ni saldo. Ambos valores se derivan de la fecha de vencimiento y del pago asociado.",
    )

    doc.add_heading("7. Límites del cierre", level=1)
    agregar_parrafo(
        doc,
        "No forman parte de este cierre las conversaciones, mensajes, WhatsApp, IA, comprobantes, OCR, "
        "tickets, reportes avanzados ni despliegue productivo. Esos componentes permanecen como módulos futuros.",
    )

    doc.add_heading("8. Lista de aceptación", level=1)
    agregar_vinetas(
        doc,
        [
            "El código está versionado y publicado en la rama main.",
            "La base de datos se crea mediante migraciones reproducibles.",
            "Las operaciones del Módulo 1 están disponibles en panel y API.",
            "Los permisos bloquean operaciones no autorizadas.",
            "Las pruebas y el análisis estático finalizan correctamente.",
            "Los pendientes están identificados y no se confunden con funcionalidades terminadas.",
        ],
    )
    return guardar(doc, "Revision_final_Modulo_1_Villafane_Wifi.docx")


def crear_manual_usuario():
    """Genera el manual operativo para administradores y empleados."""

    doc = configurar_documento("Manual de usuario", preset="compacto")
    agregar_portada(
        doc,
        "Manual de usuario",
        "Operación del panel web - Módulo 1",
        "Guía operativa",
    )

    doc.add_heading("1. Propósito del manual", level=1)
    agregar_parrafo(
        doc,
        "Este manual explica cómo utilizar las funciones disponibles del sistema: iniciar sesión, consultar el "
        "resumen, administrar clientes, planes y servicios, generar cuotas, consultar cuentas corrientes y registrar pagos.",
    )
    agregar_nota(
        doc,
        "Alcance",
        "Las funciones de WhatsApp, bot, comprobantes, OCR, tickets y reportes avanzados todavía no están disponibles.",
    )

    doc.add_heading("2. Acceso al sistema", level=1)
    doc.add_heading("Paso 1. Iniciar el servidor local", level=2)
    agregar_parrafo(doc, "Abrir PowerShell en la carpeta del proyecto y ejecutar:")
    agregar_nota(
        doc,
        "Comandos",
        ".\\.venv\\Scripts\\Activate.ps1   |   python backend\\manage.py runserver",
    )
    doc.add_heading("Paso 2. Abrir la aplicación", level=2)
    agregar_parrafo(doc, "Ingresar en http://127.0.0.1:8000/iniciar-sesion/ desde un navegador actualizado.")
    doc.add_heading("Paso 3. Autenticarse", level=2)
    agregar_vinetas(
        doc,
        [
            "Escribir el nombre de usuario entregado por el administrador.",
            "Escribir la contraseña.",
            "Seleccionar Ingresar al sistema.",
            "Si los datos son incorrectos, revisar mayúsculas y volver a intentarlo.",
        ],
    )

    doc.add_heading("3. Navegación y permisos", level=1)
    agregar_parrafo(
        doc,
        "El menú lateral muestra únicamente los módulos autorizados. Intentar abrir manualmente una dirección "
        "sin permiso produce una pantalla de acceso restringido.",
    )
    permisos = [
        ("Administrador", "Todas las operaciones del sistema."),
        ("Administración", "Clientes, planes, servicios, cuentas, pagos y reportes habilitados."),
        ("Soporte", "Consulta de clientes y servicios; soporte se ampliará con tickets."),
        ("Atención", "Consulta de clientes y cuentas; atención se ampliará con conversaciones."),
    ]
    agregar_tabla(doc, ["Perfil", "Acceso actual"], permisos, [2300, 7060])

    doc.add_heading("4. Pantalla de inicio", level=1)
    agregar_parrafo(
        doc,
        "El inicio muestra la cantidad de clientes, planes y servicios activos. Los accesos rápidos permiten "
        "registrar un cliente, buscarlo, crear un plan o consultar cuentas corrientes.",
    )

    doc.add_heading("5. Gestión de clientes", level=1)
    doc.add_heading("5.1 Buscar y consultar", level=2)
    agregar_vinetas(
        doc,
        [
            "Ingresar a Clientes.",
            "Escribir nombre, documento, teléfono/WhatsApp o localidad.",
            "Seleccionar Buscar y abrir la ficha correspondiente.",
            "Revisar datos de contacto, teléfonos y servicios contratados.",
        ],
    )
    doc.add_heading("5.2 Registrar un cliente", level=2)
    agregar_vinetas(
        doc,
        [
            "Seleccionar Nuevo cliente.",
            "Elegir tipo y número de documento.",
            "Completar nombre o razón social, tipo de cliente y dirección.",
            "Registrar al menos un teléfono válido.",
            "Completar los datos de la primera conexión y seleccionar un plan activo.",
            "Guardar y comprobar la ficha generada.",
        ],
    )
    agregar_nota(
        doc,
        "Validaciones",
        "No se admiten documentos, teléfonos, IP o MAC duplicados. El día de vencimiento debe estar entre 1 y 31.",
        tipo="advertencia",
    )
    doc.add_heading("5.3 Editar o dar de baja", level=2)
    agregar_parrafo(
        doc,
        "Desde la ficha seleccionar Editar cliente para modificar datos y teléfonos. Dar de baja solicita "
        "confirmación, inactiva el cliente y también sus conexiones, pero conserva el historial.",
    )

    doc.add_heading("6. Planes y conexiones", level=1)
    doc.add_heading("6.1 Administrar planes", level=2)
    agregar_vinetas(
        doc,
        [
            "Abrir Planes y seleccionar Nuevo plan.",
            "Completar nombre, velocidad en Mbps, precio vigente y estado.",
            "Editar un plan cuando cambien sus condiciones comerciales.",
            "Dar de baja para impedir nuevas contrataciones sin borrar el historial.",
        ],
    )
    doc.add_heading("6.2 Agregar otra conexión", level=2)
    agregar_parrafo(
        doc,
        "Desde la ficha del cliente seleccionar Agregar en Servicios contratados. Cada conexión mantiene su "
        "propio plan, dirección de instalación, vencimiento, IP, MAC y estado.",
    )

    doc.add_heading("7. Cuenta corriente", level=1)
    doc.add_heading("7.1 Consultar la situación", level=2)
    agregar_vinetas(
        doc,
        [
            "Abrir Cuenta corriente.",
            "Buscar al cliente por nombre o documento.",
            "Revisar total pendiente, deuda vencida y próximo vencimiento.",
            "Abrir la cuenta para consultar cuotas y pagos históricos.",
        ],
    )
    doc.add_heading("7.2 Crear una cuenta receptora", level=2)
    agregar_parrafo(
        doc,
        "Antes del primer pago, ingresar a Cuentas receptoras, seleccionar Nueva cuenta e indicar nombre, "
        "tipo, alias/CBU/CVU o identificación de caja, y estado.",
    )
    doc.add_heading("7.3 Generar cuotas", level=2)
    agregar_vinetas(
        doc,
        [
            "Seleccionar Generar cuotas.",
            "Confirmar el período con formato AAAA-MM.",
            "Ejecutar la generación.",
            "El sistema crea solamente las cuotas faltantes de servicios activos y evita duplicados.",
        ],
    )
    doc.add_heading("7.4 Registrar un pago", level=2)
    agregar_vinetas(
        doc,
        [
            "Abrir la cuenta corriente del cliente.",
            "Seleccionar Registrar pago.",
            "Marcar una o varias cuotas completas del mismo cliente.",
            "Elegir cuenta receptora, medio de pago, fecha y hora.",
            "Confirmar y verificar que las cuotas figuren como pagadas.",
        ],
    )
    agregar_nota(
        doc,
        "Regla de pago",
        "No existen pagos parciales. Una transferencia puede cancelar varias cuotas y servicios del mismo cliente.",
    )

    doc.add_heading("8. Estados utilizados", level=1)
    estados = [
        ("Cliente", "Activo / Inactivo", "La baja es lógica."),
        ("Servicio", "Activo / Suspendido / Inactivo", "Determina su situación operativa."),
        ("Plan", "Activo / Inactivo", "Solo los activos se asignan a nuevas conexiones."),
        ("Cuota", "Pendiente / Vencida / Pagada", "Se calcula automáticamente."),
        ("Cuenta", "Al día / Pendiente / Con deuda", "Se calcula desde las cuotas."),
    ]
    agregar_tabla(doc, ["Elemento", "Estados", "Interpretación"], estados, [1700, 3000, 4660])

    doc.add_heading("9. Mensajes frecuentes", level=1)
    mensajes = [
        ("Usuario o contraseña incorrectos", "Revisar credenciales y teclado de mayúsculas."),
        ("Acceso restringido", "El perfil no tiene permiso para esa operación."),
        ("Registro duplicado", "Buscar el cliente, teléfono, IP o MAC antes de volver a cargarlo."),
        ("Sin cuenta receptora", "Crear una cuenta receptora activa antes de registrar pagos."),
        ("Cuota ya pagada", "Revisar el historial; no debe imputarse nuevamente."),
    ]
    agregar_tabla(doc, ["Mensaje", "Acción recomendada"], mensajes, [3400, 5960])

    doc.add_heading("10. Cierre de sesión y buenas prácticas", level=1)
    agregar_vinetas(
        doc,
        [
            "Usar el botón de salida ubicado junto al usuario.",
            "No compartir credenciales entre empleados.",
            "Verificar cliente, período y cuotas antes de confirmar un pago.",
            "No eliminar registros directamente desde la base de datos.",
            "Comunicar al administrador cualquier dato duplicado o acceso incorrecto.",
        ],
    )
    return guardar(doc, "Manual_de_usuario_Modulo_1_Villafane_Wifi.docx")


def crear_documentacion_tecnica():
    """Genera la referencia técnica reproducible del backend implementado."""

    doc = configurar_documento("Documentación técnica", preset="compacto")
    agregar_portada(
        doc,
        "Documentación técnica",
        "Arquitectura, código, base de datos y operación - Módulo 1",
        "Referencia para desarrollo",
    )

    doc.add_heading("1. Visión técnica", level=1)
    agregar_parrafo(
        doc,
        "El sistema se implementa como un monolito modular con Django. Cada aplicación representa una "
        "capacidad del negocio y expone servicios de aplicación para evitar que las vistas concentren reglas. "
        "El panel web y la API REST comparten modelos, políticas de acceso y lógica transaccional.",
    )

    doc.add_heading("2. Tecnologías", level=1)
    tecnologias = [
        ("Lenguaje", "Python 3.12", "POO y reglas del negocio"),
        ("Framework", "Django 5.2", "Modelos, vistas, formularios y autenticación"),
        ("API", "Django REST Framework 3.16", "Endpoints autenticados"),
        ("Base de datos", "PostgreSQL", "Persistencia principal"),
        ("Tareas futuras", "Celery + Redis", "Automatización e integraciones"),
        ("Pruebas", "Pytest + pytest-django", "Pruebas unitarias y recorridos"),
        ("Calidad", "Ruff", "Análisis estático y formato"),
    ]
    agregar_tabla(doc, ["Capa", "Tecnología", "Responsabilidad"], tecnologias, [1900, 2800, 4660])

    doc.add_heading("3. Estructura del proyecto", level=1)
    estructura = [
        ("backend/config", "Configuración, rutas, ASGI y WSGI."),
        ("backend/apps/usuarios", "Usuario, subtipos, autenticación y permisos."),
        ("backend/apps/clientes", "Clientes, teléfonos y casos de uso."),
        ("backend/apps/servicios", "Planes y conexiones contratadas."),
        ("backend/apps/facturacion", "Cuotas, cuenta corriente y pagos."),
        ("backend/templates", "Pantallas HTML del panel."),
        ("backend/static", "CSS, logo y recursos visuales."),
        ("requirements", "Dependencias base y de desarrollo."),
        ("docs", "Convenciones, modularización y entregables."),
    ]
    agregar_tabla(doc, ["Ruta", "Contenido"], estructura, [3100, 6260])

    doc.add_heading("4. Módulos y dependencias", level=1)
    agregar_vinetas(
        doc,
        [
            "usuarios es transversal y autoriza las operaciones.",
            "clientes es propietario de Cliente y TelefonoCliente.",
            "servicios depende de clientes y administra Plan y Servicio.",
            "facturacion depende de servicios y administra Cuota, Pago y CuentaReceptora.",
            "reportes podrá leer otros módulos, pero no modificar sus datos.",
            "integraciones traducirá Meta, OCR y LLM a operaciones internas.",
        ],
    )

    doc.add_heading("5. Modelo de datos implementado", level=1)
    entidades = [
        ("Usuario", "id", "Credencial común y estado de acceso."),
        ("Empleado", "id_usuario (PK/FK)", "Subtipo con área."),
        ("Administrador", "id_usuario (PK/FK)", "Subtipo administrativo."),
        ("Cliente", "id", "Persona o empresa titular."),
        ("ClienteTelefono", "numero", "Teléfono único del cliente."),
        ("Plan", "id", "Oferta comercial y precio vigente."),
        ("Servicio", "id", "Conexión de un cliente con un plan."),
        ("CuentaReceptora", "id", "Destino habilitado para cobranzas."),
        ("Pago", "id", "Cobranza acreditada."),
        ("Cuota", "id", "Importe mensual de un servicio."),
    ]
    agregar_tabla(doc, ["Entidad", "Clave", "Responsabilidad"], entidades, [2350, 2600, 4410])

    doc.add_heading("5.1 Relaciones", level=2)
    relaciones = [
        ("Usuario - Empleado", "1:1", "Especialización con PK compartida."),
        ("Usuario - Administrador", "1:1", "Especialización con PK compartida."),
        ("Cliente - Teléfono", "1:N", "Un cliente posee varios teléfonos."),
        ("Cliente - Servicio", "1:N", "Un cliente contrata varias conexiones."),
        ("Plan - Servicio", "1:N", "Un plan se asigna a varios servicios."),
        ("Servicio - Cuota", "1:N", "Un servicio genera cuotas mensuales."),
        ("CuentaReceptora - Pago", "1:N", "Una cuenta recibe pagos."),
        ("Pago - Cuota", "1:N", "Un pago cancela varias cuotas."),
    ]
    agregar_tabla(doc, ["Relación", "Cardinalidad", "Lectura"], relaciones, [3000, 1500, 4860])

    doc.add_heading("5.2 Restricciones relevantes", level=2)
    agregar_vinetas(
        doc,
        [
            "Documento único por tipo de documento.",
            "Teléfono único en cliente_telefono.",
            "Nombre de plan único sin distinguir mayúsculas.",
            "IP y MAC únicas cuando están informadas.",
            "Día de vencimiento entre 1 y 31.",
            "Una cuota por servicio y período.",
            "Montos de cuota y pago positivos.",
            "Fecha de vencimiento posterior o igual a la emisión.",
        ],
    )

    doc.add_heading("6. Diseño orientado a objetos", level=1)
    agregar_parrafo(
        doc,
        "Los modelos representan entidades persistentes. Las clases de servicios coordinan validaciones y "
        "transacciones. Los formularios y serializadores validan entradas, mientras las vistas se ocupan de HTTP y presentación.",
    )
    servicios = [
        ("ServicioUsuarios", "Crea un usuario con exactamente un subtipo."),
        ("ServicioAutorizacion", "Selecciona la política según subtipo y área."),
        ("ServicioClientes", "Alta, búsqueda, edición y baja lógica."),
        ("ServicioPlanes", "Gestión del catálogo comercial."),
        ("ServicioContrataciones", "Altas, cambios y bajas de conexiones."),
        ("ServicioFacturacion", "Genera cuotas y registra pagos atómicos."),
        ("ServicioCuentaCorriente", "Calcula deuda e historial sin redundancia."),
    ]
    agregar_tabla(doc, ["Clase", "Responsabilidad"], servicios, [3300, 6060])

    doc.add_heading("7. Seguridad y autorización", level=1)
    agregar_parrafo(
        doc,
        "Django almacena las contraseñas con hash y utiliza sesiones. La generalización Usuario-Empleado-Administrador "
        "es total y disjunta a nivel de servicios de creación. La política de autorización rechaza perfiles faltantes o dobles.",
    )
    agregar_nota(
        doc,
        "Producción",
        "Activar DEBUG=False, HTTPS, SECURE_SSL_REDIRECT, SESSION_COOKIE_SECURE, CSRF_COOKIE_SECURE y HSTS cuando corresponda.",
        tipo="advertencia",
    )

    doc.add_heading("8. Rutas principales", level=1)
    rutas_panel = [
        ("/", "Inicio del panel"),
        ("/iniciar-sesion/", "Autenticación"),
        ("/panel/clientes/", "Clientes"),
        ("/panel/planes/", "Planes"),
        ("/panel/conexiones/", "Servicios"),
        ("/panel/cuentas/", "Cuentas corrientes"),
        ("/panel/cuentas/generar-cuotas/", "Generación mensual"),
        ("/admin/", "Administración técnica de Django"),
        ("/health/", "Comprobación de salud"),
    ]
    agregar_tabla(doc, ["Ruta", "Función"], rutas_panel, [4200, 5160])

    doc.add_heading("8.1 API REST", level=2)
    rutas_api = [
        ("/api/clientes/", "CRUD con baja lógica"),
        ("/api/planes/", "Catálogo de planes"),
        ("/api/servicios/", "Conexiones contratadas"),
        ("/api/cuotas/", "Consulta de cuotas"),
        ("/api/cuotas/generar/", "Generación de cuotas"),
        ("/api/pagos/", "Consulta y registro de pagos"),
    ]
    agregar_tabla(doc, ["Endpoint", "Responsabilidad"], rutas_api, [4200, 5160])

    doc.add_heading("9. Configuración local", level=1)
    agregar_parrafo(doc, "Variables esperadas en .env, sin incluir valores secretos:")
    agregar_vinetas(
        doc,
        [
            "DJANGO_SETTINGS_MODULE y DJANGO_SECRET_KEY.",
            "DJANGO_ALLOWED_HOSTS.",
            "DATABASE_URL para PostgreSQL.",
            "REDIS_URL para tareas futuras.",
            "Credenciales de WhatsApp, LLM y Google Cloud cuando se implementen integraciones.",
        ],
    )
    doc.add_heading("9.1 Preparación", level=2)
    agregar_vinetas(
        doc,
        [
            "Crear el entorno: python -m venv .venv.",
            "Activarlo: .\\.venv\\Scripts\\Activate.ps1.",
            "Instalar dependencias: pip install -r requirements\\dev.txt.",
            "Crear la base villafane_wifi en PostgreSQL.",
            "Aplicar migraciones: python backend\\manage.py migrate.",
            "Ejecutar: python backend\\manage.py runserver.",
        ],
    )

    doc.add_heading("10. Pruebas y calidad", level=1)
    agregar_parrafo(
        doc,
        "La batería actual contiene 50 pruebas. Las pruebas rápidas usan SQLite en memoria; el módulo de facturación "
        "también fue validado contra PostgreSQL. La configuración real continúa utilizando PostgreSQL.",
    )
    controles = [
        ("pytest -q", "Ejecuta las 50 pruebas."),
        ("ruff check backend", "Analiza estilo, errores e imports."),
        ("manage.py check", "Valida configuración de Django."),
        ("makemigrations --check --dry-run", "Detecta cambios sin migración."),
        ("manage.py check --deploy", "Advierte configuraciones inseguras de producción."),
    ]
    agregar_tabla(doc, ["Comando", "Objetivo"], controles, [4100, 5260])

    doc.add_heading("11. Despliegue previsto", level=1)
    agregar_parrafo(
        doc,
        "La aplicación está preparada para ejecutarse con Gunicorn o Uvicorn y una base PostgreSQL administrada. "
        "Railway o Render son opciones para desarrollo y demostración. Antes del despliegue deben configurarse "
        "dominio, HTTPS, variables secretas, archivos estáticos, copias de seguridad y monitoreo.",
    )

    doc.add_heading("12. Evolución prevista", level=1)
    agregar_vinetas(
        doc,
        [
            "Conversaciones y mensajes.",
            "Integración con WhatsApp y transferencia a empleados.",
            "Clasificación de intención mediante LLM externo.",
            "Comprobantes, hash de duplicados, OCR y conciliación.",
            "Tickets, notas internas y soporte técnico.",
            "Reportes, alertas y automatizaciones con Celery.",
        ],
    )

    doc.add_heading("13. Criterios de mantenimiento", level=1)
    agregar_vinetas(
        doc,
        [
            "Revisar requerimientos y DER antes de cada módulo.",
            "Agregar migraciones por cada modificación persistente.",
            "Ubicar reglas de negocio en servicios, no en plantillas.",
            "Mantener nombres y documentación propia en español.",
            "Agregar pruebas para cada regla nueva o error corregido.",
            "No publicar archivos .env, credenciales ni bases locales.",
        ],
    )
    return guardar(doc, "Documentacion_tecnica_Modulo_1_Villafane_Wifi.docx")


def principal():
    """Genera los tres entregables y muestra sus rutas."""

    rutas = [crear_revision_final(), crear_manual_usuario(), crear_documentacion_tecnica()]
    for ruta in rutas:
        print(ruta)


if __name__ == "__main__":
    principal()
