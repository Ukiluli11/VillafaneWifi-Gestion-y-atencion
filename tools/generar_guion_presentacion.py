"""Genera el guion compartido para la exposición del proyecto Villafañe Wifi."""

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from generar_documentacion_modulo1 import (
    GRIS,
    VERDE,
    VERDE_OSCURO,
    agregar_nota,
    agregar_parrafo,
    agregar_portada,
    agregar_tabla,
    agregar_vinetas,
    configurar_documento,
    configurar_fuente,
    guardar,
)


def agregar_bloque_guion(documento, numero, titulo, expositor, tiempo, apoyo, texto):
    """Agrega una intervención oral con responsable, duración y apoyo visual."""

    documento.add_heading(f"{numero}. {titulo}", level=1)

    responsable = documento.add_paragraph()
    responsable.paragraph_format.space_after = Pt(4)
    configurar_fuente(responsable.add_run(f"Habla: {expositor}"), color=VERDE, negrita=True)
    configurar_fuente(responsable.add_run(f"  |  Tiempo estimado: {tiempo}"), color=GRIS, cursiva=True)

    indicacion = documento.add_paragraph()
    indicacion.paragraph_format.space_after = Pt(7)
    configurar_fuente(indicacion.add_run(f"[Apoyo visual: {apoyo}]"), color=GRIS, cursiva=True)

    for parrafo_texto in texto:
        parrafo = documento.add_paragraph()
        parrafo.paragraph_format.left_indent = Pt(14)
        parrafo.paragraph_format.right_indent = Pt(8)
        parrafo.paragraph_format.space_after = Pt(7)
        configurar_fuente(parrafo.add_run(parrafo_texto))


def crear_guion():
    """Construye el documento completo del guion de exposición."""

    documento = configurar_documento("Guion de exposición", preset="compacto")
    agregar_portada(
        documento,
        "Guion de exposición del proyecto",
        "Presentación compartida por Ulises y Agustín - duración estimada: 15 a 17 minutos",
        "Guion para defensa y demostración",
    )

    documento.add_heading("Cómo utilizar este guion", level=1)
    agregar_parrafo(
        documento,
        "El texto está preparado para ser expresado con naturalidad, sin necesidad de leerlo de forma literal. "
        "Las indicaciones entre corchetes señalan qué documento, diagrama o pantalla conviene mostrar. La "
        "distribución busca que ambos integrantes participen de manera equilibrada.",
    )
    agregar_nota(
        documento,
        "Antes de comenzar:",
        "tener abierto el informe, los diagramas, el sistema funcionando y una cuenta de demostración. "
        "No mostrar contraseñas, archivos .env ni datos personales reales.",
        tipo="advertencia",
    )
    documento.add_heading("Distribución general", level=2)
    agregar_tabla(
        documento,
        ["Expositor", "Ejes principales", "Tiempo aproximado"],
        [
            ("Ulises", "Introducción, solución, documentación, arquitectura y primera parte de la demostración", "8 minutos"),
            ("Agustín", "Problemática, alcance, modelos de datos, segunda parte de la demostración y calidad", "8 minutos"),
        ],
        [1450, 6410, 1500],
    )

    agregar_bloque_guion(
        documento,
        "1",
        "Apertura y presentación",
        "Ulises",
        "1 minuto",
        "portada del informe o diapositiva con el nombre y el logo del proyecto",
        [
            "Buenos días. Somos Ulises Belazquez y Agustín Serrano, estudiantes de la Licenciatura en Sistemas de Información. En el marco de Seminario de Integración desarrollamos el proyecto denominado Sistema de Gestión Integral y Atención al Cliente para Villafañe Wifi.",
            "El objetivo del trabajo es aplicar de manera integrada los conocimientos de análisis, bases de datos, ingeniería de software y programación para resolver necesidades reales de una empresa proveedora de internet.",
        ],
    )

    agregar_bloque_guion(
        documento,
        "2",
        "Contexto y problemática relevada",
        "Agustín",
        "1 minuto y 30 segundos",
        "sección Contexto del proyecto y síntesis de la entrevista",
        [
            "Villafañe Wifi administra clientes, conexiones y cobranzas con información distribuida principalmente en planillas y comunicaciones por WhatsApp. Esta modalidad permite operar, pero genera tareas repetitivas y dificulta mantener una única fuente de información.",
            "Durante el relevamiento identificamos problemas como la búsqueda manual de clientes, el seguimiento de servicios y cuotas, la verificación de pagos y la atención de reclamos. También observamos que una misma persona puede tener más de una conexión y que cada servicio debe conservar sus propios datos técnicos y comerciales.",
            "A partir de esta situación definimos requisitos funcionales y no funcionales, priorizando primero la gestión interna y dejando las integraciones más complejas para módulos posteriores.",
        ],
    )

    agregar_bloque_guion(
        documento,
        "3",
        "Solución propuesta y objetivos",
        "Ulises",
        "1 minuto y 30 segundos",
        "objetivo general, objetivos específicos y esquema general del sistema",
        [
            "La solución propuesta es un sistema centralizado compuesto por un panel web para administradores y empleados y, en etapas posteriores, un canal automatizado de atención por WhatsApp. Ambos componentes compartirán la misma lógica de negocio y la misma base de datos.",
            "Nuestro objetivo general es optimizar la gestión comercial, administrativa y de atención al cliente. Como objetivos específicos buscamos centralizar los datos, administrar planes y conexiones, consultar cuentas corrientes, registrar pagos, preservar el historial, controlar los accesos y reducir errores derivados de la carga manual.",
            "La construcción se organizó de manera modular. Esto nos permite entregar una base funcional y continuar incorporando capacidades sin tener que rehacer el sistema completo.",
        ],
    )

    agregar_bloque_guion(
        documento,
        "4",
        "Alcance actual y exclusiones",
        "Agustín",
        "1 minuto y 15 segundos",
        "sección Alcance y exclusiones del informe",
        [
            "En el alcance del primer módulo se encuentran la autenticación y los permisos, la gestión de clientes, planes y conexiones, la cuenta corriente, la generación de cuotas y el registro de pagos completos.",
            "Todavía no forman parte de la versión terminada el bot de WhatsApp, la interpretación mediante inteligencia artificial, el OCR de comprobantes, la conciliación automática, los tickets y los reportes avanzados. Tampoco se contempla por el momento la integración con Mikrotik para cortar o rehabilitar automáticamente el servicio.",
            "Estas exclusiones no son olvidos: están documentadas como evolución futura para mantener un alcance realizable y verificable durante la cursada.",
        ],
    )

    agregar_bloque_guion(
        documento,
        "5",
        "Documentación y planificación elaborada",
        "Ulises",
        "1 minuto y 30 segundos",
        "índice del informe, requerimientos y diagrama de Gantt",
        [
            "Antes de programar elaboramos la documentación que guía el proyecto. El informe incluye el contexto, la elicitación, la solución propuesta, los objetivos, el alcance, los requisitos funcionales y no funcionales, la arquitectura tecnológica y la planificación.",
            "Los requerimientos se numeraron para poder relacionarlos con las tareas del Gantt, la implementación y las pruebas. El cronograma separa el análisis y la documentación inicial de los módulos funcionales, e incorpora pruebas, manuales y cierre de cada etapa.",
            "Además, preparamos una revisión final del Módulo 1, un manual de usuario y una documentación técnica. De esta forma distinguimos qué necesita conocer un usuario operativo, qué debe mantener un desarrollador y qué evidencia permite evaluar el avance.",
        ],
    )

    agregar_bloque_guion(
        documento,
        "6",
        "Diagramas y modelo de datos",
        "Agustín",
        "1 minuto y 40 segundos",
        "DER conceptual, DER lógico y diagrama de casos de uso",
        [
            "El DER conceptual representa las entidades y reglas principales del negocio. Un cliente puede contratar varios servicios; cada servicio corresponde a un plan y genera cuotas mensuales. Un pago puede cancelar varias cuotas completas del mismo cliente.",
            "Para los usuarios del sistema aplicamos una generalización: Usuario es el supertipo y Empleado y Administrador son subtipos. En el modelo lógico esta herencia se implementa mediante una clave primaria compartida, que también funciona como clave foránea.",
            "También definimos restricciones para evitar documentos, teléfonos, direcciones IP y direcciones MAC duplicados. Los saldos y estados de cuenta se calculan a partir de cuotas y pagos, evitando almacenar información redundante. Los casos de uso complementan el modelo mostrando las operaciones disponibles para cada perfil.",
        ],
    )

    agregar_bloque_guion(
        documento,
        "7",
        "Tecnologías, arquitectura y forma de trabajo",
        "Ulises",
        "1 minuto y 30 segundos",
        "diagrama de arquitectura y documentación técnica",
        [
            "El backend y el panel web están desarrollados con Python, Django y Django REST Framework. Utilizamos PostgreSQL como base de datos y organizamos el código como un monolito modular, con aplicaciones separadas para usuarios, clientes, servicios y facturación.",
            "La programación sigue un enfoque orientado a objetos. Los modelos representan las entidades persistentes y las clases de servicio concentran las reglas de negocio. Los nombres de clases, funciones, variables, comentarios y mensajes se mantienen mayormente en español para facilitar su comprensión académica y el mantenimiento.",
            "Para las etapas futuras prevemos la API oficial de WhatsApp de Meta, un servicio externo de lenguaje para interpretar intenciones, Google Cloud Vision con Tesseract como respaldo para OCR y Celery con Redis para tareas automáticas.",
        ],
    )

    agregar_bloque_guion(
        documento,
        "8",
        "Demostración: acceso y gestión comercial",
        "Ulises",
        "2 minutos",
        "sistema abierto en la pantalla de inicio de sesión",
        [
            "Ahora vamos a mostrar el funcionamiento del Módulo 1. Primero ingresamos con un usuario autorizado. El menú se adapta a los permisos del administrador o al área del empleado, por lo que no todos los perfiles pueden realizar las mismas acciones.",
            "Desde el panel podemos consultar el resumen general y acceder a Clientes. La búsqueda permite localizar registros por nombre, documento, teléfono o localidad. Al crear un cliente se registran sus datos de contacto y su primera conexión; luego pueden agregarse otras conexiones, cada una con plan, dirección de instalación, vencimiento, IP y MAC propios.",
            "También podemos administrar el catálogo de planes, modificar sus condiciones o inactivarlos sin eliminar el historial existente.",
        ],
    )

    agregar_bloque_guion(
        documento,
        "9",
        "Demostración: cuenta corriente y pagos",
        "Agustín",
        "2 minutos",
        "cuenta corriente de un cliente de demostración",
        [
            "En Cuenta corriente se observa la deuda total, la deuda vencida, el próximo vencimiento y el historial del cliente. Para cada período se generan únicamente las cuotas faltantes de los servicios activos, evitando duplicaciones.",
            "Al registrar un pago se selecciona una cuenta receptora, el medio utilizado y una o varias cuotas completas del mismo cliente. El sistema no admite pagos parciales porque esa es la regla definida durante el relevamiento.",
            "Después de confirmar, las cuotas aparecen pagadas y el saldo se recalcula automáticamente. Los registros no se eliminan físicamente: cuando corresponde se utiliza una baja lógica para conservar la trazabilidad.",
        ],
    )

    agregar_bloque_guion(
        documento,
        "10",
        "Pruebas, seguridad y estado actual",
        "Agustín",
        "1 minuto",
        "revisión final del Módulo 1 o resultado de pruebas",
        [
            "El Módulo 1 cuenta con cincuenta pruebas automatizadas que verifican modelos, servicios, permisos, panel y API. También ejecutamos análisis estático con Ruff, comprobaciones de Django y control de migraciones.",
            "La versión es adecuada para demostración académica y para continuar el desarrollo. Antes de una publicación productiva todavía debemos configurar HTTPS, desactivar el modo de depuración, utilizar secretos exclusivos del entorno y establecer copias de seguridad y monitoreo.",
        ],
    )

    agregar_bloque_guion(
        documento,
        "11",
        "Continuidad del proyecto",
        "Ulises",
        "45 segundos",
        "Gantt o lista de módulos futuros",
        [
            "La siguiente etapa incorporará conversaciones y mensajes, integración con WhatsApp, transferencia de la atención a empleados, clasificación de intenciones, comprobantes con control de duplicados y OCR, tickets de soporte y reportes.",
            "La base actual ya fue diseñada para crecer por módulos, manteniendo separados los datos, las reglas de negocio y las interfaces.",
        ],
    )

    agregar_bloque_guion(
        documento,
        "12",
        "Cierre",
        "Agustín y Ulises",
        "45 segundos",
        "pantalla inicial del sistema o portada del proyecto",
        [
            "Agustín: Como resultado, logramos transformar las necesidades relevadas en requisitos, modelos, una planificación y un primer módulo funcional con base de datos y control de accesos.",
            "Ulises: El proyecto no se limita a programar pantallas; integra análisis del negocio, diseño de datos, arquitectura, pruebas y documentación. Con esta base podemos continuar los módulos de atención y automatización de manera ordenada.",
            "Ambos: Muchas gracias. Quedamos atentos a sus preguntas.",
        ],
    )

    documento.add_heading("Preguntas posibles y responsable sugerido", level=1)
    agregar_tabla(
        documento,
        ["Pregunta", "Responde", "Idea central de la respuesta"],
        [
            ("¿Por qué eligieron Django?", "Ulises", "Experiencia previa, POO, autenticación, ORM, rapidez de desarrollo y API REST."),
            ("¿Por qué PostgreSQL?", "Agustín", "Integridad, restricciones, transacciones y adecuación a un sistema relacional."),
            ("¿Por qué un cliente puede tener varios servicios?", "Agustín", "Una persona puede contratar conexiones en distintas ubicaciones o con distintos planes."),
            ("¿Por qué no implementaron todavía WhatsApp?", "Ulises", "Se priorizó una base interna estable; la integración corresponde al siguiente módulo."),
            ("¿Cómo evitan inconsistencias?", "Agustín", "Validaciones, claves únicas, transacciones, bajas lógicas y cálculo derivado de saldos."),
            ("¿El sistema ya está listo para producción?", "Ulises", "Está listo para demostración; faltan endurecimiento de seguridad, despliegue, respaldos y monitoreo."),
        ],
        [4300, 1200, 3860],
    )

    documento.add_page_break()
    documento.add_heading("Lista de control antes de exponer", level=1)
    agregar_vinetas(
        documento,
        [
            "Confirmar que PostgreSQL y el servidor de Django estén funcionando.",
            "Preparar datos ficticios para la demostración: cliente, plan, servicio, cuotas y cuenta receptora.",
            "Abrir previamente el informe, el Gantt, los diagramas y el panel web.",
            "Ensayar las transiciones entre ambos expositores sin interrumpirse.",
            "Cronometrar una práctica completa y reducir ejemplos si supera los 17 minutos.",
            "Evitar mostrar credenciales, datos personales reales o archivos de configuración.",
        ],
    )

    cierre = documento.add_paragraph()
    cierre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cierre.paragraph_format.space_before = Pt(18)
    configurar_fuente(
        cierre.add_run("Guion preparado para la presentación del Módulo 1 - Villafañe Wifi"),
        tamano=10,
        color=VERDE_OSCURO,
        negrita=True,
    )

    return guardar(documento, "Guion_exposicion_Ulises_y_Agustin_Villafane_Wifi.docx")


if __name__ == "__main__":
    print(crear_guion())
